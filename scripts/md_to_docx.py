#!/usr/bin/env python3
"""Minimal md -> docx converter for the Sivrce tech decision doc.
Handles: # ## ### headings, tables, - and 1. lists, **bold**, `code`,
--- rules, fenced code blocks. ponytail: single-purpose, no deps beyond python-docx."""
import re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC, DST = sys.argv[1], sys.argv[2]
doc = Document()
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.9)

st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(10.5)

ACCENT = RGBColor(0x1F, 0x3B, 0x73)

def add_runs(p, text):
    # split **bold** and `code`
    for tok in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not tok: continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        else:
            p.add_run(tok)

def flush_table(rows):
    cells = [[c.strip() for c in re.split(r"(?<!\\)\|", r.strip())[1:-1]] for r in rows]
    if len(cells) > 1 and all(re.fullmatch(r":?-{2,}:?", c or "---") for c in cells[1]):
        cells.pop(1)  # separator
    if not cells: return
    t = doc.add_table(rows=len(cells), cols=len(cells[0]))
    t.style = "Table Grid"
    for i, row in enumerate(cells):
        for j, val in enumerate(row[: len(cells[0])]):
            cell = t.cell(i, j)
            p = cell.paragraphs[0]
            add_runs(p, re.sub(r"\\\|", "|", val))
            for r in p.runs:
                r.font.size = Pt(9)
                if i == 0: r.bold = True
    doc.add_paragraph()

lines = open(SRC, encoding="utf-8").read().splitlines()
i, in_code = 0, False
while i < len(lines):
    ln = lines[i]
    if ln.strip().startswith("```"):
        in_code = not in_code
        if in_code:
            buf = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i]); i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(buf)); r.font.name = "Consolas"; r.font.size = Pt(8.5)
        i += 1; continue
    if ln.startswith("|"):
        rows = []
        while i < len(lines) and lines[i].startswith("|"):
            rows.append(lines[i]); i += 1
        flush_table(rows); continue
    if ln.startswith("#"):
        m = re.match(r"(#+)\s+(.*)", ln)
        lvl, txt = len(m.group(1)), m.group(2)
        h = doc.add_heading(level=min(lvl, 4))
        add_runs(h, txt)
        for r in h.runs: r.font.color.rgb = ACCENT
    elif re.match(r"\s*[-*]\s+", ln):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, re.sub(r"^\s*[-*]\s+", "", ln))
    elif re.match(r"\s*\d+\.\s+", ln):
        p = doc.add_paragraph(style="List Number")
        add_runs(p, re.sub(r"^\s*\d+\.\s+", "", ln))
    elif ln.strip() == "---":
        pass
    elif ln.strip():
        p = doc.add_paragraph()
        add_runs(p, ln)
    i += 1

# footer page numbers via field
sec = doc.sections[0]
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)

doc.save(DST)
print("saved", DST)
